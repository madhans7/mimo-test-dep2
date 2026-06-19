from fpdf import FPDF
from docx import Document
from PIL import Image, ImageDraw, ImageFont
import os

# 1. 5-Page PDF
pdf = FPDF()
for i in range(1, 6):
    pdf.add_page()
    pdf.set_font("Arial", size=24)
    pdf.cell(200, 10, txt=f"This is Test Page {i} of 5", ln=True, align='C')
pdf.output("test_5_pages.pdf")
print("Created test_5_pages.pdf")

# 2. Word Document
doc = Document()
doc.add_heading('Mimo Test Document', 0)
doc.add_paragraph('This is a test Word document to verify LibreOffice conversion on the Pi.')
doc.save("test_doc.docx")
print("Created test_doc.docx")

# 3. Image File
img = Image.new('RGB', (800, 600), color = (73, 109, 137))
d = ImageDraw.Draw(img)
# Just simple text
d.text((100,250), "Test Image for Mimo Kiosk", fill=(255,255,0))
img.save('test_image.jpg')
print("Created test_image.jpg")
